from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Any, List, Optional
import io

from dependencies import get_current_user_id

router = APIRouter(prefix="/api/export", tags=["export"])


# ── Request schemas ────────────────────────────────────────────────────────────

class ExportLine(BaseModel):
    line: int
    text: str

class ScriptXlsxRequest(BaseModel):
    title: Optional[str] = "script"
    rows: List[ExportLine]

class ScriptDocxRequest(BaseModel):
    title: Optional[str] = "script"
    text: str

class ScriptGoogleDocsRequest(BaseModel):
    title: Optional[str] = "script"
    text: str
    rendered_html: Optional[str] = None
    folder_id: Optional[str] = None
    google_access_token: str
    docs_blocks: Optional[List[dict]] = None

class ReportRow(BaseModel):
    category: str
    content: str
    line: str

class ReportXlsxRequest(BaseModel):
    title: Optional[str] = "script_report"
    columns: List[str]
    rows: List[ReportRow]

class ReportDocxRequest(BaseModel):
    title: Optional[str] = "script_report"
    doc_title: Optional[str] = "統計報表"
    columns: List[str]
    rows: List[ReportRow]


# ── Helpers ────────────────────────────────────────────────────────────────────

def _safe_filename(name: str, ext: str) -> str:
    safe = "".join(c for c in (name or "script") if c.isalnum() or c in "-_ ")
    return f"{safe or 'script'}.{ext}"


def _build_xlsx(columns: List[str], data_rows: list) -> bytes:
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
    except ImportError:
        raise HTTPException(status_code=500, detail="openpyxl not installed")

    wb = openpyxl.Workbook()
    ws = wb.active

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(fill_type="solid", fgColor="111827")

    for col_idx, col_name in enumerate(columns, start=1):
        cell = ws.cell(row=1, column=col_idx, value=col_name)
        cell.font = header_font
        cell.fill = header_fill

    for row_idx, row in enumerate(data_rows, start=2):
        for col_idx, value in enumerate(row, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


def _build_docx_from_html(title: str, text: str) -> bytes:
    try:
        from docx import Document
        from docx.shared import RGBColor
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed")

    doc = Document()
    doc.add_heading(title, level=1)
    for line in text.split("\n"):
        doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _parse_css_flags(style_attr: str) -> dict:
    style = (style_attr or "").lower()
    color = None
    if "color:" in style:
        try:
            part = style.split("color:", 1)[1].split(";", 1)[0].strip()
            if part.startswith("#") and len(part) == 7:
                color = part[1:].upper()
            elif part.startswith("rgb(") and part.endswith(")"):
                nums = [int(x.strip()) for x in part[4:-1].split(",")[:3]]
                if len(nums) == 3:
                    color = "".join(f"{max(0, min(255, n)):02X}" for n in nums)
        except Exception:
            color = None
    return {
        "bold": ("font-weight:bold" in style) or ("font-weight:700" in style) or ("font-weight:800" in style) or ("font-weight:900" in style),
        "italic": "font-style:italic" in style,
        "underline": "text-decoration:underline" in style or "text-decoration-line:underline" in style,
        "color": color,
    }


def _add_styled_runs(paragraph, node, parent_style):
    from bs4 import NavigableString, Tag
    from docx.shared import RGBColor

    if isinstance(node, NavigableString):
        text = str(node)
        if not text:
            return
        run = paragraph.add_run(text)
        run.bold = bool(parent_style.get("bold"))
        run.italic = bool(parent_style.get("italic"))
        run.underline = bool(parent_style.get("underline"))
        color = parent_style.get("color")
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        return

    if not isinstance(node, Tag):
        return

    next_style = dict(parent_style)
    tag_name = (node.name or "").lower()
    if tag_name in {"strong", "b"}:
        next_style["bold"] = True
    if tag_name in {"em", "i"}:
        next_style["italic"] = True
    if tag_name in {"u"}:
        next_style["underline"] = True

    inline_style = _parse_css_flags(node.attrs.get("style", ""))
    if inline_style.get("bold"):
        next_style["bold"] = True
    if inline_style.get("italic"):
        next_style["italic"] = True
    if inline_style.get("underline"):
        next_style["underline"] = True
    if inline_style.get("color"):
        next_style["color"] = inline_style["color"]

    for child in node.children:
        _add_styled_runs(paragraph, child, next_style)


def _build_docx_from_rendered_html(title: str, rendered_html: str, fallback_text: str = "") -> bytes:
    try:
        from docx import Document
        from bs4 import BeautifulSoup
    except ImportError:
        return _build_docx_from_html(title, fallback_text)

    doc = Document()
    doc.add_heading(title, level=1)

    html = rendered_html or ""
    if not html.strip():
        for line in (fallback_text or "").split("\n"):
            doc.add_paragraph(line)
    else:
        soup = BeautifulSoup(html, "html.parser")
        lines = soup.select(".script-line")
        if not lines:
            lines = soup.find_all(["p", "div", "li"])

        if lines:
            for line in lines:
                para = doc.add_paragraph()
                _add_styled_runs(para, line, {"bold": False, "italic": False, "underline": False, "color": None})
        else:
            plain = soup.get_text("\n")
            for line in plain.split("\n"):
                doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _create_google_doc_via_docx_import(
    title: str,
    text: str,
    rendered_html: Optional[str],
    google_access_token: str,
    folder_id: Optional[str],
) -> dict:
    try:
        from googleapiclient.discovery import build
        from google.oauth2.credentials import Credentials
        from googleapiclient.http import MediaIoBaseUpload
    except ImportError:
        raise HTTPException(status_code=500, detail="google-api-python-client not installed")

    creds = Credentials(token=google_access_token)
    drive_service = build("drive", "v3", credentials=creds)
    docx_bytes = _build_docx_from_rendered_html(title or "Script", rendered_html or "", text or "")
    media = MediaIoBaseUpload(
        io.BytesIO(docx_bytes),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        resumable=False,
    )

    clean_folder_id = (folder_id or "").strip()
    metadata = {
        "name": title or "Script",
        "mimeType": "application/vnd.google-apps.document",
    }
    if clean_folder_id:
        metadata["parents"] = [clean_folder_id]

    created = drive_service.files().create(
        body=metadata,
        media_body=media,
        fields="id",
        supportsAllDrives=True,
    ).execute()
    document_id = created.get("id")
    if not document_id:
        raise HTTPException(status_code=500, detail="Failed to create Google Doc from DOCX import")

    return {
        "documentId": document_id,
        "documentUrl": f"https://docs.google.com/document/d/{document_id}/edit",
        "exportMode": "docx_import_fallback",
    }


def _hex_to_rgb(color_hex: str) -> Optional[dict]:
    raw = str(color_hex or "").strip().lstrip("#")
    if len(raw) != 6:
        return None
    try:
        r = int(raw[0:2], 16) / 255.0
        g = int(raw[2:4], 16) / 255.0
        b = int(raw[4:6], 16) / 255.0
        return {"red": r, "green": g, "blue": b}
    except Exception:
        return None


def _css_color_to_rgb(raw_color: str) -> Optional[dict]:
    raw = str(raw_color or "").strip()
    if not raw:
        return None
    if raw.startswith("#"):
        short = raw.lstrip("#")
        if len(short) == 3:
            raw = f"#{short[0]}{short[0]}{short[1]}{short[1]}{short[2]}{short[2]}"
        return _hex_to_rgb(raw)
    if raw.lower().startswith("rgb(") and raw.endswith(")"):
        try:
            chunks = [part.strip() for part in raw[4:-1].split(",")]
            if len(chunks) != 3:
                return None
            r = max(0, min(255, int(float(chunks[0])))) / 255.0
            g = max(0, min(255, int(float(chunks[1])))) / 255.0
            b = max(0, min(255, int(float(chunks[2])))) / 255.0
            return {"red": r, "green": g, "blue": b}
        except Exception:
            return None
    return None


def _css_px_to_pt_dimension(value: Any) -> Optional[dict]:
    try:
        px = float(value)
    except Exception:
        return None
    if px < 0:
        return None
    return {"magnitude": round(px * 0.75, 2), "unit": "PT"}


def _build_padding_style(style_obj: dict) -> tuple[dict, List[str]]:
    style: dict = {}
    fields: List[str] = []
    for source_key, docs_key in [
        ("paddingTop", "paddingTop"),
        ("paddingRight", "paddingRight"),
        ("paddingBottom", "paddingBottom"),
        ("paddingLeft", "paddingLeft"),
    ]:
        dim = _css_px_to_pt_dimension(style_obj.get(source_key))
        if dim:
            style[docs_key] = dim
            fields.append(docs_key)
    return style, fields


def _build_column_width_requests(table_start_index: int, table_layout: Optional[dict], n_cols: int) -> List[dict]:
    if not isinstance(table_layout, dict):
        return []
    widths = table_layout.get("columnWidths")
    if not isinstance(widths, list) or len(widths) < n_cols:
        return []
    numeric_widths = []
    for raw in widths[:n_cols]:
        try:
            width = float(raw)
        except Exception:
            width = 0
        numeric_widths.append(width if width > 0 else 0)
    total = sum(numeric_widths)
    if total <= 0:
        return []

    target_table_width_pt = 468.0
    requests = []
    for col_idx, width in enumerate(numeric_widths):
        width_pt = max(5.0, round((width / total) * target_table_width_pt, 2))
        requests.append({
            "updateTableColumnProperties": {
                "tableStartLocation": {"index": table_start_index},
                "columnIndices": [col_idx],
                "tableColumnProperties": {
                    "widthType": "FIXED_WIDTH",
                    "width": {"magnitude": width_pt, "unit": "PT"},
                },
                "fields": "width,widthType",
            }
        })
    return requests


def _create_google_doc_from_blocks(
    title: str,
    docs_blocks: List[dict],
    google_access_token: str,
    folder_id: Optional[str],
) -> dict:
    try:
        from googleapiclient.discovery import build
        from google.oauth2.credentials import Credentials
    except ImportError:
        raise HTTPException(status_code=500, detail="google-api-python-client not installed")

    creds = Credentials(token=google_access_token)
    docs_service = build("docs", "v1", credentials=creds)
    drive_service = build("drive", "v3", credentials=creds)

    doc = docs_service.documents().create(body={"title": title or "Script"}).execute()
    document_id = doc.get("documentId")
    if not document_id:
        raise HTTPException(status_code=500, detail="Failed to create Google Doc")

    insert_chunks: List[str] = []
    style_requests: List[dict] = []
    cursor = 1
    for block in docs_blocks or []:
      runs = block.get("runs") if isinstance(block, dict) else None
      if not isinstance(runs, list) or len(runs) == 0:
          insert_chunks.append("\n")
          cursor += 1
          continue

      for run in runs:
          if not isinstance(run, dict):
              continue
          text = str(run.get("text", ""))
          if not text:
              continue
          start = cursor
          end = start + len(text)
          insert_chunks.append(text)
          fields: List[str] = []
          style: dict = {}
          if run.get("bold"):
              style["bold"] = True
              fields.append("bold")
          if run.get("italic"):
              style["italic"] = True
              fields.append("italic")
          if run.get("underline"):
              style["underline"] = True
              fields.append("underline")
          rgb = _hex_to_rgb(str(run.get("color", "")))
          if rgb:
              style["foregroundColor"] = {"color": {"rgbColor": rgb}}
              fields.append("foregroundColor")
          if fields:
              style_requests.append({
                  "updateTextStyle": {
                      "range": {"startIndex": start, "endIndex": end},
                      "textStyle": style,
                      "fields": ",".join(fields),
                  }
              })
          cursor = end
      insert_chunks.append("\n")
      cursor += 1

    all_text = "".join(insert_chunks) or "\n"
    requests = [{"insertText": {"location": {"index": 1}, "text": all_text}}]
    requests.extend(style_requests)
    docs_service.documents().batchUpdate(
        documentId=document_id,
        body={"requests": requests},
    ).execute()

    clean_folder_id = (folder_id or "").strip()
    if clean_folder_id:
        file_meta = drive_service.files().get(fileId=document_id, fields="parents").execute()
        previous_parents = ",".join(file_meta.get("parents", []))
        drive_service.files().update(
            fileId=document_id,
            addParents=clean_folder_id,
            removeParents=previous_parents or None,
            fields="id, parents",
            supportsAllDrives=True,
        ).execute()

    return {
        "documentId": document_id,
        "documentUrl": f"https://docs.google.com/document/d/{document_id}/edit",
        "exportMode": "docs_blocks",
    }


def _build_report_docx(doc_title: str, columns: List[str], rows: List[ReportRow]) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed")

    doc = Document()
    doc.add_heading(doc_title, level=1)
    doc.add_paragraph(f"共 {len(rows)} 筆資料")

    table = doc.add_table(rows=1 + len(rows), cols=len(columns))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = col
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True

    for row_idx, row in enumerate(rows, start=1):
        cells = table.rows[row_idx].cells
        cells[0].text = row.category
        cells[1].text = row.content
        cells[2].text = str(row.line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Routes ─────────────────────────────────────────────────────────────────────

@router.post("/xlsx")
async def export_script_xlsx(
    req: ScriptXlsxRequest,
    _uid: str = Depends(get_current_user_id),
):
    columns = ["行號", "內容"]
    data_rows = [[r.line, r.text] for r in req.rows]
    xlsx_bytes = _build_xlsx(columns, data_rows)
    filename = _safe_filename(req.title, "xlsx")
    return StreamingResponse(
        io.BytesIO(xlsx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/docx")
async def export_script_docx(
    req: ScriptDocxRequest,
    _uid: str = Depends(get_current_user_id),
):
    docx_bytes = _build_docx_from_html(req.title or "Script", req.text)
    filename = _safe_filename(req.title, "docx")
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/google-docs")
async def export_script_google_docs(
    req: ScriptGoogleDocsRequest,
    _uid: str = Depends(get_current_user_id),
):
    try:
        if isinstance(req.docs_blocks, list) and len(req.docs_blocks) > 0:
            return _create_google_doc_from_blocks(
                title=req.title or "Script",
                docs_blocks=req.docs_blocks,
                google_access_token=req.google_access_token,
                folder_id=req.folder_id,
            )
        return _create_google_doc_via_docx_import(
            title=req.title or "Script",
            text=req.text,
            rendered_html=req.rendered_html,
            google_access_token=req.google_access_token,
            folder_id=req.folder_id,
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Google Docs export failed: {e}")


def _create_google_doc_table(
    title: str,
    columns: List[str],
    rows: List[List[str]],
    cell_styles: Optional[List[List[Optional[dict]]]],
    cell_runs: Optional[List[List[List[dict]]]],
    table_layout: Optional[dict],
    google_access_token: str,
    folder_id: Optional[str],
) -> dict:
    """Create a Google Doc with a multi-column table using the Docs API."""
    try:
        from googleapiclient.discovery import build
        from google.oauth2.credentials import Credentials
    except ImportError:
        raise HTTPException(status_code=500, detail="google-api-python-client not installed")

    creds = Credentials(token=google_access_token)
    docs_service = build("docs", "v1", credentials=creds)
    drive_service = build("drive", "v3", credentials=creds)

    doc = docs_service.documents().create(body={"title": title or "Script"}).execute()
    document_id = doc.get("documentId")
    if not document_id:
        raise HTTPException(status_code=500, detail="Failed to create Google Doc")

    n_cols = len(columns)
    n_data_rows = len(rows)
    total_rows = 1 + n_data_rows  # header + data

    # Step 1: insert table at index 1
    docs_service.documents().batchUpdate(
        documentId=document_id,
        body={"requests": [
            {"insertTable": {"rows": total_rows, "columns": n_cols, "location": {"index": 1}}}
        ]},
    ).execute()

    # Step 2: re-fetch doc to get accurate cell indices
    doc_state = docs_service.documents().get(documentId=document_id).execute()
    body_content = doc_state.get("body", {}).get("content", [])

    # Collect all table cell startIndex values (paragraph inside each cell)
    cell_indices: List[int] = []
    table_start_index = 1
    for elem in body_content:
        table = elem.get("table")
        if not table:
            continue
        table_start_index = elem.get("startIndex") or table_start_index
        for table_row in table.get("tableRows", []):
            for cell in table_row.get("tableCells", []):
                for content_item in cell.get("content", []):
                    if "paragraph" in content_item:
                        cell_indices.append(content_item["startIndex"])
                        break  # only first paragraph per cell

    # Step 3: build insertText + bold-header requests
    # Must insert in REVERSE order to preserve indices
    all_cells = [columns[i] for i in range(n_cols)]  # header row
    for row in rows:
        for col_idx in range(n_cols):
            all_cells.append(row[col_idx] if col_idx < len(row) else "")

    text_requests = []
    table_layout_requests = _build_column_width_requests(table_start_index, table_layout, n_cols)
    cell_style_requests = []
    if isinstance(table_layout, dict) and isinstance(table_layout.get("defaultCellStyle"), dict):
        default_cell_style, default_cell_fields = _build_padding_style(table_layout["defaultCellStyle"])
        if default_cell_fields:
            cell_style_requests.append({
                "updateTableCellStyle": {
                    "tableStartLocation": {"index": table_start_index},
                    "tableCellStyle": default_cell_style,
                    "fields": ",".join(default_cell_fields),
                }
            })

    for i, (idx, text) in enumerate(zip(cell_indices, all_cells)):
        row_idx = i // n_cols
        col_idx = i % n_cols

        source_runs = None
        if row_idx > 0 and isinstance(cell_runs, list) and (row_idx - 1) < len(cell_runs):
            run_row = cell_runs[row_idx - 1]
            if isinstance(run_row, list) and col_idx < len(run_row):
                source_runs = run_row[col_idx]

        if row_idx == 0:
            normalized_runs = [{"text": str(text or ""), "bold": True}]
        elif isinstance(source_runs, list) and len(source_runs) > 0:
            normalized_runs = [run for run in source_runs if isinstance(run, dict)]
        else:
            normalized_runs = [{"text": str(text or "")}]

        safe_text = "".join(str(run.get("text", "")) for run in normalized_runs)
        if safe_text:
            text_requests.append({
                "insertText": {"location": {"index": idx}, "text": safe_text}
            })
            offset = 0
            for run in normalized_runs:
                run_text = str(run.get("text", ""))
                if not run_text:
                    continue
                text_fields: List[str] = []
                text_style: dict = {}
                if run.get("bold") is True:
                    text_style["bold"] = True
                    text_fields.append("bold")
                if run.get("italic") is True:
                    text_style["italic"] = True
                    text_fields.append("italic")
                if run.get("underline") is True:
                    text_style["underline"] = True
                    text_fields.append("underline")
                text_rgb = _css_color_to_rgb(str(run.get("color", "")))
                if text_rgb:
                    text_style["foregroundColor"] = {"color": {"rgbColor": text_rgb}}
                    text_fields.append("foregroundColor")
                if text_fields:
                    text_requests.append({
                        "updateTextStyle": {
                            "range": {"startIndex": idx + offset, "endIndex": idx + offset + len(run_text)},
                            "textStyle": text_style,
                            "fields": ",".join(text_fields),
                        }
                    })
                offset += len(run_text)

        style_obj = None
        if row_idx > 0 and isinstance(cell_styles, list) and (row_idx - 1) < len(cell_styles):
            style_row = cell_styles[row_idx - 1]
            if isinstance(style_row, list) and col_idx < len(style_row):
                style_obj = style_row[col_idx]

        if isinstance(style_obj, dict):
            table_cell_style, table_cell_fields = _build_padding_style(style_obj)
            bg_rgb = _css_color_to_rgb(str(style_obj.get("backgroundColor", "")))
            if bg_rgb:
                table_cell_style["backgroundColor"] = {"color": {"rgbColor": bg_rgb}}
                table_cell_fields.append("backgroundColor")
            if table_cell_fields:
                cell_style_requests.append({
                    "updateTableCellStyle": {
                        "tableRange": {
                            "tableCellLocation": {
                                "tableStartLocation": {"index": table_start_index},
                                "rowIndex": row_idx,
                                "columnIndex": col_idx,
                            },
                            "rowSpan": 1,
                            "columnSpan": 1,
                        },
                        "tableCellStyle": table_cell_style,
                        "fields": ",".join(table_cell_fields),
                    }
                })

    # Apply per-cell text requests from the end of the document toward the start
    # so each cell's insertion index remains valid while its style requests stay
    # adjacent to the inserted text.
    grouped_text_requests = []
    current_group = []
    for request in text_requests:
        if "insertText" in request and current_group:
            grouped_text_requests.append(current_group)
            current_group = [request]
        else:
            current_group.append(request)
    if current_group:
        grouped_text_requests.append(current_group)
    text_requests = [request for group in reversed(grouped_text_requests) for request in group]
    cell_style_requests.reverse()

    if table_layout_requests or text_requests or cell_style_requests:
        docs_service.documents().batchUpdate(
            documentId=document_id,
            body={"requests": table_layout_requests + text_requests + cell_style_requests},
        ).execute()

    clean_folder_id = (folder_id or "").strip()
    if clean_folder_id:
        file_meta = drive_service.files().get(fileId=document_id, fields="parents").execute()
        previous_parents = ",".join(file_meta.get("parents", []))
        drive_service.files().update(
            fileId=document_id,
            addParents=clean_folder_id,
            removeParents=previous_parents or None,
            fields="id, parents",
            supportsAllDrives=True,
        ).execute()

    return {
        "documentId": document_id,
        "documentUrl": f"https://docs.google.com/document/d/{document_id}/edit",
        "exportMode": "table_v2",
    }


class GoogleDocsTableV2Request(BaseModel):
    title: Optional[str] = "script"
    google_access_token: str
    folder_id: Optional[str] = None
    columns: List[str]
    rows: List[List[str]]
    cell_styles: Optional[List[List[Optional[dict]]]] = None
    cell_runs: Optional[List[List[List[dict]]]] = None
    table_layout: Optional[dict] = None


@router.post("/google-docs/v2")
async def export_google_docs_table_v2(
    req: GoogleDocsTableV2Request,
    _uid: str = Depends(get_current_user_id),
):
    try:
        return _create_google_doc_table(
            title=req.title or "Script",
            columns=req.columns,
            rows=req.rows,
            cell_styles=req.cell_styles,
            cell_runs=req.cell_runs,
            table_layout=req.table_layout,
            google_access_token=req.google_access_token,
            folder_id=req.folder_id,
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Google Docs table export failed: {e}")


class TableV2Request(BaseModel):
    title: Optional[str] = "script"
    doc_title: Optional[str] = None
    columns: List[str]
    rows: List[List[str]]


def _build_table_v2_docx(doc_title: str, columns: List[str], rows: List[List[str]]) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed")

    doc = Document()
    doc.add_heading(doc_title, level=1)
    doc.add_paragraph(f"共 {len(rows)} 筆資料")

    n_cols = len(columns)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = col
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True

    for row_idx, row in enumerate(rows, start=1):
        cells = table.rows[row_idx].cells
        for col_idx in range(n_cols):
            cells[col_idx].text = str(row[col_idx]) if col_idx < len(row) else ""

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@router.post("/report/v2/xlsx")
async def export_report_v2_xlsx(
    req: TableV2Request,
    _uid: str = Depends(get_current_user_id),
):
    xlsx_bytes = _build_xlsx(req.columns, req.rows)
    filename = _safe_filename(req.title, "xlsx")
    return StreamingResponse(
        io.BytesIO(xlsx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/report/v2/docx")
async def export_report_v2_docx(
    req: TableV2Request,
    _uid: str = Depends(get_current_user_id),
):
    docx_bytes = _build_table_v2_docx(req.doc_title or req.title or "劇本", req.columns, req.rows)
    filename = _safe_filename(req.title, "docx")
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/report/xlsx")
async def export_report_xlsx(
    req: ReportXlsxRequest,
    _uid: str = Depends(get_current_user_id),
):
    data_rows = [[r.category, r.content, r.line] for r in req.rows]
    xlsx_bytes = _build_xlsx(req.columns, data_rows)
    filename = _safe_filename(req.title, "xlsx")
    return StreamingResponse(
        io.BytesIO(xlsx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/report/docx")
async def export_report_docx(
    req: ReportDocxRequest,
    _uid: str = Depends(get_current_user_id),
):
    docx_bytes = _build_report_docx(req.doc_title or req.title or "報表", req.columns, req.rows)
    filename = _safe_filename(req.title, "docx")
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
